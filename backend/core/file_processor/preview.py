# file_processor/preview.py
import os
import io
import base64
from PIL import Image, ImageDraw, ImageFont
import PyPDF2
from docx import Document
import pptx
from django.conf import settings

class PreviewGenerator:
    """Generate previews for uploaded files"""
    
    @staticmethod
    def generate_preview(file, file_name):
        """Generate preview based on file type"""
        ext = os.path.splitext(file_name)[1][1:].lower()
        
        if ext == 'pdf':
            return PreviewGenerator._preview_pdf(file)
        elif ext in ['jpg', 'jpeg', 'png']:
            return PreviewGenerator._preview_image(file)
        elif ext in ['docx', 'doc']:
            return PreviewGenerator._preview_docx(file)
        elif ext == 'pptx':
            return PreviewGenerator._preview_pptx(file)
        else:
            return PreviewGenerator._preview_text(file)
    
    @staticmethod
    def _preview_pdf(file):
        """Generate PDF preview"""
        try:
            reader = PyPDF2.PdfReader(file)
            total_pages = len(reader.pages)
            
            # Get first page text
            first_page = reader.pages[0]
            text = first_page.extract_text()[:200] if first_page.extract_text() else "No text extractable"
            
            return {
                'type': 'pdf',
                'pages': total_pages,
                'preview': text,
                'thumbnail': None
            }
        except Exception as e:
            return {'type': 'pdf', 'error': str(e)}
    
    @staticmethod
    def _preview_image(file):
        """Generate image preview"""
        try:
            img = Image.open(file)
            
            # Create thumbnail
            img.thumbnail((200, 200))
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            thumbnail = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            return {
                'type': 'image',
                'thumbnail': f'data:image/png;base64,{thumbnail}',
                'size': img.size
            }
        except Exception as e:
            return {'type': 'image', 'error': str(e)}
    
    @staticmethod
    def _preview_docx(file):
        """Generate DOCX preview"""
        try:
            doc = Document(file)
            text = '\n'.join([p.text for p in doc.paragraphs[:10]])
            
            return {
                'type': 'document',
                'preview': text[:500] + ('...' if len(text) > 500 else ''),
                'paragraphs': len(doc.paragraphs)
            }
        except Exception as e:
            return {'type': 'document', 'error': str(e)}
    
    @staticmethod
    def _preview_pptx(file):
        """Generate PPTX preview"""
        try:
            prs = pptx.Presentation(file)
            slides = len(prs.slides)
            
            return {
                'type': 'presentation',
                'slides': slides,
                'preview': f'{slides} slides in this presentation'
            }
        except Exception as e:
            return {'type': 'presentation', 'error': str(e)}
    
    @staticmethod
    def _preview_text(file):
        """Generate text file preview"""
        try:
            content = file.read(1000).decode('utf-8', errors='ignore')
            file.seek(0)
            
            return {
                'type': 'text',
                'preview': content + ('...' if len(content) >= 1000 else '')
            }
        except Exception as e:
            return {'type': 'text', 'error': str(e)}
